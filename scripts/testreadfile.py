import docx


def readfile():
    pass
    val = docx.Document("/home/lixingyu/桌面/nogamenolife1.docx")
    for tmp in range(10):
        print(val.paragraphs[tmp].text)
    print(len(val.paragraphs))


if __name__ == "__main__":
    readfile()
