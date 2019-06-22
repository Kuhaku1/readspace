from django.shortcuts import render
from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.urls import reverse
from django.views.decorators.cache import cache_page
from .models import Book, BookRanking, Press, Family, Subsection, ThisLightNovelIsTerrific
from userspace.models import User, Usercollection, Usercomment
from tool.login import is_login
from tool.login import login_verification
from tool.readverify import have_permission_to_read
from django.utils.safestring import mark_safe

from .tests import testsclearrank

import ReadBooks.settings
import os

import docx  # python docx模块

# @cache_page(60*10)


def index(request):
    """主页（设置了缓存）"""
    context = {}
    # 获取精品推荐
    book_recommend = BookRanking.userManager.order_by('-collection')[0:6]
    booklist = []
    for tmp in book_recommend:
        # print(tmp.book)
        booklist.append(tmp.book)
    context['bookRecommend'] = booklist
    # 获得 最新 & 热门
    book_popularity = BookRanking.userManager.order_by('-popularity')[0:3]
    popularity_id = []
    for tmp in book_popularity:
        popularity_id.append(tmp.book_id)
    book_popularity = Book.userManager.filter(
        pk__in=popularity_id).order_by('-lastUpdated')
    context['bookPopularity'] = book_popularity

    bookrank2019 = ThisLightNovelIsTerrific.userManager.filter(
        years="15").order_by('ranking')[0:3]  # 15-->2019  14-->2018
    context["bookrank2019"] = bookrank2019

    return render(request, 'mainsite/index.html', context)


def cke(request):
    return render(request, 'mainsite/index.html')


def get_book_info(request, bookid):
    book = Book.userManager.filter(pk=bookid).first()
    context = dict()
    context['book'] = book
    context['bookauthor'] = book.author_set.all()
    context['bookstate'] = book.get_state_display()
    family = book.family_set.all()
    familylist = list()
    for tmp in family:
        familylist.append(tmp.get_name_display())

    subsections = Subsection.userManager.filter(
        book=book).order_by("bookNumber").all()
    # print(subsections[-1].bookNumber)
    tempsub = list()
    lastbook = ''
    for tmp1 in subsections:
        subsectionsitem = {}
        subsectionsitem['bookNumber'] = tmp1.bookNumber
        subsectionsitem['name'] = tmp1.name
        subsectionsitem['id'] = tmp1.id
        lastbook = tmp1
        tempsub.append(subsectionsitem)
    if lastbook and lastbook.updateDay:
        # color: #524e4e;
        updateinfo = "{}更新了第<b style='{}' >&nbsp;{}&nbsp;</b>卷".format(
            lastbook.updateDay.strftime("%Y-%m-%w"),
            "color:#FF0000",
            lastbook.bookNumber)
    else:
        updateinfo = "暂无数据"
    context['updateinfo'] = mark_safe(updateinfo)
    context['subsections'] = tempsub

    context['bookfamily'] = familylist
    username = is_login(request)
    context['bookid'] = bookid
    if not username:
        context['iflogin'] = False
        context['iflike'] = False
        context['ifcollection'] = False
        context['ifcoin'] = False
    else:
        context['iflogin'] = username
        collection = Usercollection.UserCollectionManager.filter(
            user__username=username, book=book).first()
        if collection:
            context['ifcollection'] = True
        else:
            context['ifcollection'] = False

        like = Usercollection.UserLikesManager.filter(
            user__username=username, book=book).first()
        if like:
            context['iflike'] = True
        else:
            context['iflike'] = False

        coin = Usercollection.UserCoinManager.filter(
            user__username=username, book=book).first()
        if coin:
            context['ifcoin'] = True
        else:
            context['ifcoin'] = False
    comments = Usercomment.userManager.filter(book=book).all()
    context['comments'] = comments

    return render(request, 'mainsite/bookinfo.html', context)


@login_verification
@have_permission_to_read
def get_book_context(request, sid):
    pass
    subsection = Subsection.userManager.filter(
        pk=sid).first()
    # print(subsection)
    if not subsection:
        return HttpResponseRedirect(reverse('mainsite:index'))

    filepath = os.path.join(ReadBooks.settings.MEDIA_ROOT,
                            subsection.physicalFileAddress.name)
    filecontext = docx.Document(filepath)

    centextlist = list()

    # 暂时实行全读取
    for tmp in filecontext.paragraphs:
        centextlist.append(tmp.text)

    # 下面留作分页加载
    # filelenth = len(filecontext.paragraphs)
    # if filelenth > 50:
        # filelenth = 50
    # for tmp in range(filelenth):
    #     centextlist.append(filecontext.paragraphs[tmp].text)

    context = dict()
    context['subsection'] = subsection
    context["centextlist"] = centextlist
    context["bookid"] = "1"
    return render(request, "mainsite/bookread.html", context)


def sceen(request):
    """检索页"""
    books = Book.userManager.order_by('id')[0:18]
    context = dict()
    context['book'] = books
    return render(request, 'mainsite/screen.html', context)


def search(request):

    return render(request, 'mainsite/search.html')


def novelrank(request):
    context = dict()
    bookrank2019 = ThisLightNovelIsTerrific.userManager.filter(
        years="15").order_by('ranking')[0:3]  # 15-->2019  14-->2018
    context["bookrank2019"] = bookrank2019
    bookrank2018 = ThisLightNovelIsTerrific.userManager.filter(
        years="14").order_by('ranking')[0:3]  # 15-->2019  14-->2018
    context["bookrank2018"] = bookrank2018
    bookrank2017 = ThisLightNovelIsTerrific.userManager.filter(
        years="13").order_by('ranking')[0:3]  # 15-->2019  14-->2018
    context["bookrank2017"] = bookrank2017
    return render(request, 'mainsite/novelrank.html', context)


def novelrankyear(request, year):
    yearindex = ""
    for tmp in ThisLightNovelIsTerrific.year_choices:
        if year == int(tmp[1]):
            yearindex = tmp[0]
    context = dict()
    if yearindex:
        bookrank = ThisLightNovelIsTerrific.userManager.filter(
            years=yearindex).order_by('ranking').all()  # 15-->2019  14-->2018
        context["bookrankyear"] = bookrank
    else:
        context["bookrankyear"] = ""
    return render(request, 'mainsite/novelrankyear.html', context)


def clear(request):
    testsclearrank()
    return HttpResponse("ok")
